import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.models.schemas import OptimizedResume


class ResumeGeneratorService:
    """Generate downloadable PDF and DOCX resumes."""

    BULLET_PREFIX = re.compile(r"^[\u2022\-*•]\s*")

    def generate(self, resume: OptimizedResume, fmt: str) -> tuple[bytes, str, str]:
        if fmt == "docx":
            content = self._generate_docx(resume)
            return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
        content = self._generate_pdf(resume)
        return content, "application/pdf", "pdf"

    def _add_header(self, doc: Document, resume: OptimizedResume) -> None:
        if resume.name:
            name = doc.add_paragraph(resume.name)
            name.runs[0].bold = True
            name.runs[0].font.size = Pt(16)
            name.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact_parts = []
        if resume.email:
            contact_parts.append(resume.email)
        if resume.phone:
            contact_parts.append(resume.phone)
        if resume.linkedin:
            contact_parts.append(f"linkedin.com/in/{resume.linkedin}")

        if contact_parts:
            contact = doc.add_paragraph(" | ".join(contact_parts))
            contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact.runs[0].font.size = Pt(10)

    def _add_header_pdf(self, story: list, resume: OptimizedResume, title_style, body_style) -> None:
        if resume.name:
            story.append(Paragraph(f"<b>{resume.name}</b>", title_style))
        contact_parts = []
        if resume.email:
            contact_parts.append(resume.email)
        if resume.phone:
            contact_parts.append(resume.phone)
        if resume.linkedin:
            contact_parts.append(f"linkedin.com/in/{resume.linkedin}")
        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), body_style))
            story.append(Spacer(1, 8))

    def _generate_docx(self, resume: OptimizedResume) -> bytes:
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        self._add_header(doc, resume)

        if resume.summary:
            doc.add_heading("Professional Summary", level=2)
            doc.add_paragraph(resume.summary)

        if resume.skills_section:
            doc.add_heading("Skills", level=2)
            for line in resume.skills_section.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())
        elif resume.skills:
            doc.add_heading("Skills", level=2)
            doc.add_paragraph(", ".join(resume.skills))

        if resume.experience:
            doc.add_heading("Experience", level=2)
            for exp in resume.experience:
                header_parts = [p for p in [exp.title, exp.company] if p]
                if header_parts:
                    p = doc.add_paragraph(" | ".join(header_parts))
                    p.runs[0].bold = True
                if exp.duration:
                    date_p = doc.add_paragraph(exp.duration)
                    date_p.runs[0].italic = True
                if exp.description:
                    for line in exp.description.split("\n"):
                        line = self.BULLET_PREFIX.sub("", line.strip())
                        if line:
                            doc.add_paragraph(line, style="List Bullet")

        if resume.education:
            doc.add_heading("Education", level=2)
            for edu in resume.education:
                doc.add_paragraph(edu, style="List Bullet")

        if resume.certifications:
            doc.add_heading("Certifications", level=2)
            for cert in resume.certifications:
                doc.add_paragraph(cert, style="List Bullet")

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _generate_pdf(self, resume: OptimizedResume) -> bytes:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "NameTitle",
            parent=styles["Heading1"],
            fontSize=16,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
        section_style = ParagraphStyle(
            "SectionHeading",
            parent=styles["Heading2"],
            fontSize=12,
            spaceBefore=14,
            spaceAfter=6,
            textColor="#1a2332",
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )
        bullet_style = ParagraphStyle(
            "Bullet",
            parent=body_style,
            leftIndent=18,
            bulletIndent=6,
        )

        story: list = []
        self._add_header_pdf(story, resume, title_style, body_style)

        if resume.summary:
            story.append(Paragraph("Professional Summary", section_style))
            story.append(Paragraph(resume.summary, body_style))
            story.append(Spacer(1, 6))

        if resume.skills_section:
            story.append(Paragraph("Skills", section_style))
            for line in resume.skills_section.split("\n"):
                if line.strip():
                    story.append(Paragraph(line.strip(), body_style))
            story.append(Spacer(1, 6))
        elif resume.skills:
            story.append(Paragraph("Skills", section_style))
            story.append(Paragraph(", ".join(resume.skills), body_style))
            story.append(Spacer(1, 6))

        if resume.experience:
            story.append(Paragraph("Experience", section_style))
            for exp in resume.experience:
                header_parts = [p for p in [exp.title, exp.company] if p]
                if header_parts:
                    story.append(Paragraph("<b>" + " | ".join(header_parts) + "</b>", body_style))
                if exp.duration:
                    story.append(Paragraph(f"<i>{exp.duration}</i>", body_style))
                if exp.description:
                    for line in exp.description.split("\n"):
                        line = self.BULLET_PREFIX.sub("", line.strip())
                        if line:
                            story.append(Paragraph(f"• {line}", bullet_style))
                story.append(Spacer(1, 4))

        if resume.education:
            story.append(Paragraph("Education", section_style))
            for edu in resume.education:
                story.append(Paragraph(f"• {edu}", bullet_style))

        if resume.certifications:
            story.append(Paragraph("Certifications", section_style))
            for cert in resume.certifications:
                story.append(Paragraph(f"• {cert}", bullet_style))

        if not story:
            story.append(
                Paragraph("Optimized Resume", ParagraphStyle("Title", alignment=TA_CENTER, fontSize=14))
            )

        doc.build(story)
        return buffer.getvalue()
